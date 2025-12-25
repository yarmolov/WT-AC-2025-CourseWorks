#!/usr/bin/env python3
"""Generate course-work assignment .docx files from a BrSTU template.

What it does
- Reads students from `students/students.csv`.
- Reads variants (01-40) from `Курсовые_работы_Веб_Технологии_Варианты_01-40.md`.
- Produces one .docx per student by replacing specific text fragments in the template.

Design notes
- The provided template docx doesn't contain explicit placeholders like {NAME}.
  It contains fixed labels (e.g. "1. Тема проекта") and empty table cells.
- Editing WordprocessingML reliably requires a docx library.
  We use `python-docx` to edit table cells while preserving formatting.

Assumptions (can be adjusted)
- We fill these fields if found in the template:
  - Project topic: "Вариант XX — <Название>" (from variants.md)
  - Task statement: paste the whole variant block (pitch + MVP + API + Data + acceptance + bonuses)
  - Student name / group: if a row in template contains the words "группа"/"ФИО", we'll try to fill neighboring cells.

If your template uses different wording, run with --dry-run to see what cells were detected.
"""

from __future__ import annotations

import argparse
import csv
import dataclasses
import datetime as dt
import os
import re
import sys
from pathlib import Path
from typing import Iterable, List, Optional, Tuple


@dataclasses.dataclass(frozen=True)
class Student:
    variant: Optional[int]
    group: str
    number: str
    sub: str
    name: str
    name_latin: str
    directory: str
    github: str


@dataclasses.dataclass(frozen=True)
class Variant:
    number: int
    title_line: str  # e.g. "Вариант 01 — ..."
    body_markdown: str


def _read_students(csv_path: Path) -> List[Student]:
    with csv_path.open("r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        students: List[Student] = []
        for row in reader:
            v = (row.get("Вариант") or "").strip()
            variant_num = int(v) if v.isdigit() else None
            students.append(
                Student(
                    variant=variant_num,
                    group=(row.get("Group") or "").strip(),
                    number=(row.get("№") or "").strip(),
                    sub=(row.get("sub") or "").strip(),
                    name=(row.get("Name") or "").strip(),
                    name_latin=(row.get("NameLatin") or "").strip(),
                    directory=(row.get("Directory") or "").strip(),
                    github=(row.get("Github Username") or "").strip(),
                )
            )
    # The file contains duplicates (likely repeated blocks). Keep first unique by (group, number, name).
    seen = set()
    uniq: List[Student] = []
    for s in students:
        k = (s.group, s.number, s.name)
        if k in seen:
            continue
        seen.add(k)
        uniq.append(s)
    return uniq


_VARIANT_HEADER_RE = re.compile(r"^##\s+Вариант\s+(\d+)\s+—\s+(.+?)\s*$", re.M)


def _read_variants(md_path: Path, max_variant: int = 40) -> dict[int, Variant]:
    text = md_path.read_text(encoding="utf-8")

    matches = list(_VARIANT_HEADER_RE.finditer(text))
    variants: dict[int, Variant] = {}

    for i, m in enumerate(matches):
        num = int(m.group(1))
        if num > max_variant:
            continue

        start = m.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        block = text[start:end].strip()

        title_line = f"Вариант {num:02d} — {m.group(2).strip()}"
        variants[num] = Variant(number=num, title_line=title_line, body_markdown=block)

    if not variants:
        raise RuntimeError(f"No variants found in {md_path}")

    return variants


def _ensure_python_docx():
    try:
        import docx  # noqa: F401
    except Exception:
        raise RuntimeError(
            "Missing dependency: python-docx. Install it with: pip install python-docx"
        )


def _set_cell_text(cell, text: str) -> None:
    # python-docx: setting cell.text clears runs; acceptable for our use.
    cell.text = text


def _find_cells_containing(doc, needle: str) -> List[Tuple[int, int, int]]:
    """Return list of (table_idx, row_idx, col_idx) whose cell text contains needle (case-insensitive)."""
    needle_l = needle.lower()
    hits: List[Tuple[int, int, int]] = []
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if needle_l in (cell.text or "").lower():
                    hits.append((ti, ri, ci))
    return hits


def _fill_doc_template(template_path: Path, out_path: Path, student: Student, variant: Variant, issued_date: Optional[str], dry_run: bool) -> None:
    _ensure_python_docx()
    from docx import Document  # type: ignore

    doc = Document(str(template_path))

    # Strategy:
    # 1) Fill project topic: the cell that contains "1. Тема проекта" -> fill the next cell to the right if exists.
    # 2) Fill task statement: the cell that contains "1.Постановка задачи" -> fill the next cell to the right if exists.
    # 3) Fill issue date: the cell containing "Дата выдачи" -> fill the next cell to the right.

    actions = []

    def fill_next_to_label(label_substr: str, value: str) -> None:
        hits = _find_cells_containing(doc, label_substr)
        if not hits:
            actions.append(f"label-not-found: {label_substr}")
            return
        # Use first occurrence
        ti, ri, ci = hits[0]
        tbl = doc.tables[ti]
        row = tbl.rows[ri]
        if ci + 1 >= len(row.cells):
            actions.append(f"no-right-cell: {label_substr}")
            return
        target = row.cells[ci + 1]
        if dry_run:
            actions.append(f"would-fill [{label_substr}] -> (table {ti} row {ri} col {ci+1})")
        else:
            _set_cell_text(target, value)
            actions.append(f"filled [{label_substr}] -> (table {ti} row {ri} col {ci+1})")

    topic = f"{variant.title_line}"
    fill_next_to_label("Тема проекта", topic)

    # Full task text: keep markdown-ish bulleting; Word will keep it as plain text.
    task_text = variant.body_markdown
    fill_next_to_label("Постановка задачи", task_text)

    if issued_date:
        fill_next_to_label("Дата выдачи", issued_date)

    # Try to fill student name / group if template contains such cues.
    # We look for cells containing these words and fill right neighbor.
    fill_next_to_label("ФИО", student.name)
    fill_next_to_label("груп", student.group)

    if dry_run:
        # Don't write output, just report
        print(f"[dry-run] {student.name} ({student.group}) variant {variant.number:02d}: {', '.join(actions)}")
        return

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main(argv: Optional[List[str]] = None) -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", default="./BrSTU/Офицальны документы/2_5346014968074496969.docx")
    ap.add_argument("--students", default="./students/students.csv")
    ap.add_argument("--variants", default="./Курсовые_работы_Веб_Технологии_Варианты_01-40.md")
    ap.add_argument("--out-dir", default="./tmp/assignments_docx")
    ap.add_argument("--issued-date", default=dt.date.today().isoformat())
    ap.add_argument("--max-variant", type=int, default=40)
    ap.add_argument(
        "--only-variant",
        type=int,
        default=None,
        help="Generate docs only for students with this variant number",
    )
    ap.add_argument(
        "--only-student",
        default=None,
        help="Generate docs only for one student (substring match against NameLatin or Name)",
    )
    ap.add_argument("--dry-run", action="store_true")
    args = ap.parse_args(argv)

    # Workspace layout here is not a single monorepo; we assume:
    # - this script lives under WT-AC-2025-CourseWorks/
    # - the BrSTU template lives in the workspace root (/home/nand/BrSTU/...)
    course_root = Path(__file__).resolve().parents[1]  # WT-AC-2025-CourseWorks/
    workspace_root = course_root.parent  # /home/nand/

    template_path = (workspace_root / args.template).resolve()
    students_path = (course_root / args.students).resolve()
    variants_path = (course_root / args.variants).resolve()
    out_dir = (course_root / args.out_dir).resolve()

    if not template_path.exists():
        print(f"Template not found: {template_path}", file=sys.stderr)
        return 2
    if not students_path.exists():
        print(f"Students CSV not found: {students_path}", file=sys.stderr)
        return 2
    if not variants_path.exists():
        print(f"Variants file not found: {variants_path}", file=sys.stderr)
        return 2

    students = _read_students(students_path)
    variants = _read_variants(variants_path, max_variant=args.max_variant)

    missing = [s for s in students if not s.variant]
    if missing:
        print(f"Warning: {len(missing)} students have no variant assigned (blank in CSV). They will be skipped.")

    if args.only_variant is not None:
        students = [s for s in students if s.variant == args.only_variant]

    if args.only_student:
        needle = args.only_student.strip().lower()
        students = [
            s
            for s in students
            if needle in (s.name_latin or "").lower() or needle in (s.name or "").lower()
        ]

    for s in students:
        if not s.variant:
            continue
        if s.variant not in variants:
            print(f"Warning: variant {s.variant} not found in variants file; skipping {s.name}")
            continue
        v = variants[s.variant]

        safe_name = re.sub(r"[^A-Za-zА-Яа-я0-9._-]+", "_", s.name_latin or s.name)
        out_path = out_dir / f"{s.group}_{s.number}_{safe_name}_variant_{s.variant:02d}.docx"

        _fill_doc_template(
            template_path=template_path,
            out_path=out_path,
            student=s,
            variant=v,
            issued_date=args.issued_date,
            dry_run=args.dry_run,
        )

    if not args.dry_run:
        print(f"Done. Output folder: {out_dir}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
