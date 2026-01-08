"""CLI checker for IT normocontrol requirements (short checklist).

This script is intended as a lightweight alternative to running pytest.
It validates a single .docx file using the hybrid approach:
- OOXML (ZIP + XML) for strict page setup and low-level formatting checks
- python-docx for high-level structure checks

Default target: tests/ПЗ.docx

Exit codes:
- 0: no errors (warnings allowed)
- 1: at least one error
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class ItNormocontrolConfig:
    """Configuration for IT short checklist checks.

    IMPORTANT: Values must be sourced from
    `scripts/standards_verification/standars_control_it_short.md`.
    """

    margins_left_mm: float
    margins_right_mm: float
    margins_top_mm: float
    margins_bottom_mm: float

    page_width_mm: float
    page_height_mm: float

    main_font_name: str
    main_font_size_pt: float
    inline_objects_font_size_pt: float

    first_line_indent_cm: float
    line_spacing_expected: float

    required_sections_in_order: list[str]


def _read_text_file(path: Path) -> str:
    """Read a UTF-8 text file."""

    return path.read_text(encoding="utf-8")


def _parse_float_ru(value: str) -> float:
    """Parse a float that may use a comma as decimal separator."""

    return float(value.strip().replace(",", "."))


def load_it_normocontrol_config(standards_md_path: Path) -> ItNormocontrolConfig:
    """Load IT normocontrol requirements from the markdown checklist.

    The repository contains multiple standards; for the IT profile we treat
    `standars_control_it_short.md` as the single source of truth.

    Args:
        standards_md_path: Path to `standars_control_it_short.md`.

    Returns:
        Parsed configuration.

    Raises:
        ValueError: If required values cannot be parsed.
    """

    text = _read_text_file(standards_md_path)

    # 1) Margins
    # Example: "Поля (мм): левое 23, правое 10, верхнее 20, нижнее 15."
    margins_match = re.search(
        r"Поля\s*\(мм\)\s*:\s*левое\s*(\d+(?:[\.,]\d+)?)\s*,\s*"
        r"правое\s*(\d+(?:[\.,]\d+)?)\s*,\s*"
        r"верхнее\s*(\d+(?:[\.,]\d+)?)\s*,\s*"
        r"нижнее\s*(\d+(?:[\.,]\d+)?)",
        text,
        flags=re.IGNORECASE,
    )
    if not margins_match:
        raise ValueError("Не удалось распарсить поля страницы из чек-листа")

    margins_left_mm = _parse_float_ru(margins_match.group(1))
    margins_right_mm = _parse_float_ru(margins_match.group(2))
    margins_top_mm = _parse_float_ru(margins_match.group(3))
    margins_bottom_mm = _parse_float_ru(margins_match.group(4))

    # 2) Font and line spacing
    # Example: "Шрифт: Times New Roman 14 pt; межстрочный интервал 1.0."
    font_match = re.search(
        r"Шрифт\s*:\s*([A-Za-z ]+?)\s*(\d+(?:[\.,]\d+)?)\s*pt\s*;\s*"
        r"межстрочный\s+интервал\s*(\d+(?:[\.,]\d+)?)",
        text,
        flags=re.IGNORECASE,
    )
    if not font_match:
        raise ValueError("Не удалось распарсить шрифт/интервал из чек-листа")

    main_font_name = font_match.group(1).strip()
    main_font_size_pt = _parse_float_ru(font_match.group(2))
    line_spacing_expected = _parse_float_ru(font_match.group(3))

    # 3) Font size inside tables/captions/figures
    # Example: "Внутри таблиц/подрисуночных подписей/на рисунках: 12 pt."
    inline_objects_match = re.search(
        r"Внутри\s+таблиц/подрисуночных\s+подписей/на\s+рисунках\s*:\s*"
        r"(\d+(?:[\.,]\d+)?)\s*pt",
        text,
        flags=re.IGNORECASE,
    )
    if not inline_objects_match:
        raise ValueError("Не удалось распарсить кегль для таблиц/подписей/рисунков")
    inline_objects_font_size_pt = _parse_float_ru(inline_objects_match.group(1))

    # 4) Paragraph first-line indent
    # Example: "Абзац: 12,5 мм."
    indent_match = re.search(
        r"Абзац\s*:\s*(\d+(?:[\.,]\d+)?)\s*мм",
        text,
        flags=re.IGNORECASE,
    )
    if not indent_match:
        raise ValueError("Не удалось распарсить абзацный отступ")
    first_line_indent_mm = _parse_float_ru(indent_match.group(1))
    first_line_indent_cm = first_line_indent_mm / 10.0

    # 5) Required document structure order
    # In the checklist it is provided as a numbered list.
    structure_block = re.search(
        r"##\s*2\)\s*Структура\s+пояснительной\s+записки(.*?)(?:\n##\s*3\)|\Z)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if not structure_block:
        raise ValueError("Не удалось найти блок структуры документа")

    required_sections_in_order: list[str] = []
    for line in structure_block.group(1).splitlines():
        item_match = re.match(r"\s*\d+\)\s*(.+?)\s*$", line)
        if item_match:
            required_sections_in_order.append(item_match.group(1).strip())

    if not required_sections_in_order:
        raise ValueError("Не удалось распарсить список разделов (порядок)")

    # Page size is implied by "Формат: A4".
    # Keep it explicit in config for checks.
    return ItNormocontrolConfig(
        margins_left_mm=margins_left_mm,
        margins_right_mm=margins_right_mm,
        margins_top_mm=margins_top_mm,
        margins_bottom_mm=margins_bottom_mm,
        page_width_mm=210,
        page_height_mm=297,
        main_font_name=main_font_name,
        main_font_size_pt=main_font_size_pt,
        inline_objects_font_size_pt=inline_objects_font_size_pt,
        first_line_indent_cm=first_line_indent_cm,
        line_spacing_expected=line_spacing_expected,
        required_sections_in_order=required_sections_in_order,
    )


def _ensure_tests_helpers_on_syspath(repo_root: Path) -> None:
    """Ensure the repository root is on sys.path.

    When executing this script directly, Python puts the script directory on
    `sys.path[0]`, not the repository root. Adding repo root enables imports
    like `tests.helpers.*`.
    """

    if str(repo_root) not in sys.path:
        sys.path.insert(0, str(repo_root))


def _resolve_repo_root() -> Path:
    """Resolve repository root from script location."""

    # scripts/standards_verification/check_it_docx.py -> repo root is 2 levels up
    return Path(__file__).resolve().parents[2]


def _extract_all_text(doc: Document) -> str:
    """Extract joined plain text from a Word document."""

    return "\n".join(p.text for p in doc.paragraphs)


def _find_section_positions(text: str, section_titles: list[str]) -> dict[str, int]:
    """Find first occurrence positions of section titles in text.

    Returns:
        A dict of title -> index in text (0-based). Missing titles are omitted.
    """

    lower_text = text.lower()
    positions: dict[str, int] = {}

    for title in section_titles:
        index = lower_text.find(title.lower())
        if index != -1:
            positions[title] = index

    return positions


def _check_page_setup(doc_name: str, doc_xml, report, config: ItNormocontrolConfig) -> None:
    """Check page size and margins using OOXML."""

    from tests.helpers.ooxml_utils import (
        get_page_margins,
        get_page_size,
        mm_to_twips,
        twips_to_mm,
    )

    margins = get_page_margins(doc_xml)
    if not margins:
        report.add_issue(
            doc_name,
            "page_setup",
            "error",
            "Поля страницы не найдены",
            expected="Поля должны быть заданы",
            actual="Поля отсутствуют",
            location="Разметка страницы → Поля",
        )
    else:
        expected = {
            "left": config.margins_left_mm,
            "right": config.margins_right_mm,
            "top": config.margins_top_mm,
            "bottom": config.margins_bottom_mm,
        }
        tolerance_mm = 1.5
        tolerance_twips = mm_to_twips(tolerance_mm)

        for key, expected_mm in expected.items():
            if key not in margins:
                report.add_issue(
                    doc_name,
                    "page_setup",
                    "error",
                    f"Поле '{key}' не задано",
                    expected=f"{expected_mm} мм",
                    actual="не задано",
                    location="Разметка страницы → Поля",
                )
                continue

            actual_twips = margins[key]
            expected_twips = mm_to_twips(expected_mm)
            diff_twips = abs(actual_twips - expected_twips)

            if diff_twips > tolerance_twips:
                report.add_issue(
                    doc_name,
                    "page_setup",
                    "error",
                    f"Некорректное поле '{key}'",
                    expected=f"{expected_mm} мм",
                    actual=f"{twips_to_mm(actual_twips):.1f} мм",
                    location="Разметка страницы → Поля → Настраиваемые поля",
                )

    page_size = get_page_size(doc_xml)
    if not page_size:
        report.add_issue(
            doc_name,
            "page_setup",
            "warning",
            "Размер страницы не найден",
            expected="A4 (210×297 мм)",
            actual="не найден",
        )
    else:
        a4_width_twips = mm_to_twips(config.page_width_mm)
        a4_height_twips = mm_to_twips(config.page_height_mm)
        tolerance = mm_to_twips(5)

        width_diff = abs(page_size["width"] - a4_width_twips)
        height_diff = abs(page_size["height"] - a4_height_twips)

        if width_diff > tolerance or height_diff > tolerance:
            report.add_issue(
                doc_name,
                "page_setup",
                "warning",
                "Размер страницы не соответствует A4",
                expected="210×297 мм",
                actual=f"{twips_to_mm(page_size['width']):.0f}×{twips_to_mm(page_size['height']):.0f} мм",
            )


def _check_paragraph_formatting(doc_name: str, doc_xml, report, config: ItNormocontrolConfig) -> None:
    """Check indentation and line spacing using OOXML (best-effort)."""

    from tests.helpers.ooxml_utils import (
        NS,
        cm_to_twips,
        get_paragraph_properties,
        twips_to_cm,
    )

    paragraphs = doc_xml.xpath(".//w:p", namespaces=NS)

    # Indent: 12.5 mm (1.25 cm)
    expected_indent = cm_to_twips(config.first_line_indent_cm)
    tolerance = cm_to_twips(0.1)  # 1mm

    invalid_indents: list[float] = []
    for p in paragraphs:
        props = get_paragraph_properties(p)
        first_line_raw = props.get("ind", {}).get("firstLine")
        if not first_line_raw:
            continue

        try:
            first_line = int(round(float(first_line_raw)))
        except (TypeError, ValueError):
            continue

        if abs(first_line - expected_indent) > tolerance:
            invalid_indents.append(twips_to_cm(first_line))

    if invalid_indents:
        examples = ", ".join(f"{cm:.2f} см" for cm in invalid_indents[:5])
        report.add_issue(
            doc_name,
            "paragraphs",
            "warning",
            f"Найдены некорректные отступы первой строки ({len(invalid_indents)} шт.)",
            expected=f"{config.first_line_indent_cm:.2f} см",
            actual=examples,
        )

    # Line spacing: 1.0 usually corresponds to w:spacing line=240 with lineRule=auto
    paragraphs_with_spacing = doc_xml.xpath(".//w:p[w:pPr/w:spacing]", namespaces=NS)

    invalid_spacing = 0
    for p in paragraphs_with_spacing:
        props = get_paragraph_properties(p)
        spacing = props.get("spacing")
        if not spacing:
            continue

        line = spacing.get("line")
        line_rule = spacing.get("lineRule")
        if not line or line_rule != "auto":
            continue

        try:
            line_val = int(line)
        except (TypeError, ValueError):
            continue

        # 240 = single, 360 = 1.5, 480 = double
        if not (220 <= line_val <= 260):
            invalid_spacing += 1

    if paragraphs_with_spacing:
        ratio = invalid_spacing / len(paragraphs_with_spacing)
        if ratio > 0.8:
            report.add_issue(
                doc_name,
                "paragraphs",
                "warning",
                "Много параграфов с явно заданным некорректным интервалом",
                expected="1.0 (одинарный)",
                actual=f"{invalid_spacing} из {len(paragraphs_with_spacing)}",
            )


def _check_fonts(doc_name: str, doc_xml, report, config: ItNormocontrolConfig) -> None:
    """Check that explicit font settings use Times New Roman and sizes 14/12pt."""

    from tests.helpers.ooxml_utils import (
        NS,
        get_run_properties,
        pt_to_half_points,
    )

    runs = doc_xml.xpath(".//w:r", namespaces=NS)
    fonts_used: set[str] = set()
    sizes: list[int] = []

    for run in runs[:250]:
        props = get_run_properties(run)

        r_fonts = props.get("rFonts")
        if r_fonts:
            for key in ("ascii", "hAnsi", "cs"):
                font_name = r_fonts.get(key)
                if font_name:
                    fonts_used.add(font_name)

        if "sz" in props:
            sizes.append(int(props["sz"]))

    if fonts_used and config.main_font_name not in fonts_used:
        report.add_issue(
            doc_name,
            "fonts",
            "error",
            "Times New Roman не найден среди явно заданных шрифтов",
            expected=config.main_font_name,
            actual=", ".join(sorted(fonts_used))[:200],
        )

    if sizes:
        size_main = pt_to_half_points(config.main_font_size_pt)
        size_table = pt_to_half_points(config.inline_objects_font_size_pt)

        allowed = {size_main, size_table}
        nonstandard = [s for s in sizes if s not in allowed]
        ratio = len(nonstandard) / len(sizes)

        if ratio > 0.5:
            report.add_issue(
                doc_name,
                "fonts",
                "warning",
                "Много runs с нестандартным явно заданным размером шрифта",
                expected=(
                    f"{int(config.main_font_size_pt)}pt (основной) или "
                    f"{int(config.inline_objects_font_size_pt)}pt (таблицы/подписи/рисунки)"
                ),
                actual=f"{len(nonstandard)} из {len(sizes)} (пример: {nonstandard[:5]})",
            )


def _check_page_numbering(docx_path: Path, doc_name: str, report) -> None:
    """Check presence of PAGE field in any header XML (best-effort, no render)."""

    import zipfile

    has_page_field = False

    with zipfile.ZipFile(docx_path, "r") as archive:
        header_files = [name for name in archive.namelist() if name.startswith("word/header") and name.endswith(".xml")]
        for header in header_files:
            content = archive.read(header)
            # A robust XML parse is possible, but this heuristic is enough for a quick check.
            # PAGE field usually appears as instrText containing 'PAGE'.
            if b"PAGE" in content.upper():
                has_page_field = True
                break

    if not header_files:
        report.add_issue(
            doc_name,
            "pagination",
            "warning",
            "Колонтитулы не найдены (header*.xml отсутствуют) — не удалось проверить нумерацию страниц",
        )
        return

    if not has_page_field:
        report.add_issue(
            doc_name,
            "pagination",
            "warning",
            "Не найдено поле PAGE в колонтитулах (не удалось подтвердить нумерацию страниц)",
            expected="Поле PAGE в правом верхнем углу",
            actual="PAGE не найден",
        )


def _check_structure(doc_name: str, doc: Document, report) -> None:
    """Check required sections and their order using plain text search.

    The exact list/order is sourced from the IT checklist markdown.
    """

    required_in_order = list(getattr(report, "_required_sections_in_order", []))
    if not required_in_order:
        # Fallback: keep previous behavior if config wiring is missing.
        required_in_order = [
            "Задание",
            "Реферат",
            "Оглавление",
            "Введение",
            "Заключение",
            "Список использованных источников",
            "Приложения",
        ]

    text = _extract_all_text(doc)
    positions = _find_section_positions(text, required_in_order)

    missing = [title for title in required_in_order if title not in positions]
    if missing:
        report.add_issue(
            doc_name,
            "structure",
            "error",
            "Не найдены обязательные разделы",
            expected=", ".join(required_in_order),
            actual=", ".join(missing),
        )
        return

    ordered_titles = sorted(positions.items(), key=lambda item: item[1])
    ordered_names = [name for name, _ in ordered_titles]
    if ordered_names != required_in_order:
        report.add_issue(
            doc_name,
            "structure",
            "warning",
            "Порядок разделов отличается от рекомендуемого",
            expected=" → ".join(required_in_order),
            actual=" → ".join(ordered_names),
        )


def _check_references(doc_name: str, doc: Document, report) -> None:
    """Check that bracketed references exist and sources section looks numbered."""

    text = _extract_all_text(doc)

    citations = re.findall(r"\[(\d+)\]", text)
    if not citations:
        report.add_issue(
            doc_name,
            "references",
            "warning",
            "Не найдены ссылки вида [N] в тексте",
            expected="Ссылки в квадратных скобках (например: [8])",
            actual="не найдено",
        )
        return

    max_citation = max(int(n) for n in citations)

    # Heuristic: find the sources section and count numbered lines after it.
    lower_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    sources_index = None
    for i, line in enumerate(lower_lines):
        if line.lower() == "список использованных источников":
            sources_index = i
            break

    if sources_index is None:
        report.add_issue(
            doc_name,
            "references",
            "error",
            "Есть ссылки [N], но не найден раздел 'Список использованных источников'",
            expected="Раздел со списком источников",
            actual="не найден",
        )
        return

    sources_lines = lower_lines[sources_index + 1 : sources_index + 80]
    numbered = [line for line in sources_lines if re.match(r"^\d+\s+", line)]

    if not numbered:
        report.add_issue(
            doc_name,
            "references",
            "warning",
            "В разделе источников не найдены строки, начинающиеся с номера",
            expected="Нумерация арабскими цифрами без точки (например: 1 ...)",
            actual="не найдено",
        )
        return

    if max_citation > len(numbered):
        report.add_issue(
            doc_name,
            "references",
            "warning",
            "Максимальный номер ссылки больше числа найденных источников",
            expected=f"Источников ≥ {max_citation}",
            actual=f"Найдено источников (эвристика): {len(numbered)}",
        )


def _check_captions(doc_name: str, doc: Document, report) -> None:
    """Check basic caption formats for figures and tables (best-effort)."""

    figure_re = re.compile(r"^рисунок\s+\d+(?:\.\d+)?\s*[—–-]\s+.+$", re.IGNORECASE)
    table_re = re.compile(r"^таблица\s+\d+(?:\.\d+)?\s*[—–-]\s+.+$", re.IGNORECASE)

    bad_figures = 0
    bad_tables = 0

    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue

        if line.lower().startswith("рисунок"):
            if not figure_re.match(line) or line.endswith("."):
                bad_figures += 1

        if line.lower().startswith("таблица"):
            if not table_re.match(line) or line.endswith("."):
                bad_tables += 1

    if bad_figures:
        report.add_issue(
            doc_name,
            "figures",
            "warning",
            "Найдены подписи рисунков с нарушением формата",
            expected="Рисунок N – Название (без точки в конце)",
            actual=f"проблемных подписей: {bad_figures}",
        )

    if bad_tables:
        report.add_issue(
            doc_name,
            "tables",
            "warning",
            "Найдены названия таблиц с нарушением формата",
            expected="Таблица N – Название (без точки в конце)",
            actual=f"проблемных названий: {bad_tables}",
        )


def check_it_docx(docx_path: Path, report_dir: Path) -> int:
    """Run IT short checklist checks and write a markdown report.

    Args:
        docx_path: Path to a .docx file.
        report_dir: Directory where a markdown report will be saved.

    Returns:
        Exit code (0 if no errors, 1 otherwise).
    """

    repo_root = _resolve_repo_root()
    _ensure_tests_helpers_on_syspath(repo_root)

    from tests.helpers.ooxml_utils import get_document_xml
    from tests.helpers.report import NormocontrolReport

    standards_md = repo_root / "scripts" / "standards_verification" / "standars_control_it_short.md"
    config = load_it_normocontrol_config(standards_md)

    report = NormocontrolReport()
    # Pass required sections through the report instance without changing its public API.
    # (This keeps changes localized to this script.)
    setattr(report, "_required_sections_in_order", config.required_sections_in_order)
    doc_name = docx_path.name
    report.add_document(doc_name)

    doc_xml = get_document_xml(docx_path)
    doc = Document(docx_path)

    _check_page_setup(doc_name, doc_xml, report, config)
    _check_paragraph_formatting(doc_name, doc_xml, report, config)
    _check_fonts(doc_name, doc_xml, report, config)
    _check_page_numbering(docx_path, doc_name, report)
    _check_structure(doc_name, doc, report)
    _check_references(doc_name, doc, report)
    _check_captions(doc_name, doc, report)

    report_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_path = report_dir / f"it_normocontrol_report_{timestamp}.md"
    report.to_markdown(report_path)

    summary = report.generate_summary()
    print(f"✓ Report: {report_path}")
    print(f"Checked: {summary['total_documents']} document(s)")
    print(f"Issues: {summary['total_issues']} (errors={summary['errors']}, warnings={summary['warnings']})")

    return 1 if report.has_errors() else 0


def main() -> int:
    """CLI entrypoint."""

    repo_root = _resolve_repo_root()
    default_docx = repo_root / "tests" / "ПЗ.docx"

    docx_path = Path(sys.argv[1]) if len(sys.argv) >= 2 else default_docx
    report_dir = repo_root / "normocontrol_reports"

    if not docx_path.exists():
        print(f"ERROR: File not found: {docx_path}")
        return 1

    if docx_path.suffix.lower() != ".docx":
        print(f"ERROR: Expected .docx file: {docx_path}")
        return 1

    return check_it_docx(docx_path, report_dir)


if __name__ == "__main__":
    raise SystemExit(main())
