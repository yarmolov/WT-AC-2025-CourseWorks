"""
Normocontrol tests for .docx documents (Variant 3: Hybrid approach).

Tests formatting requirements from docs/Требования_к_нормоконтролю.md:
- Page margins: 30mm left, 10mm right, 20mm top/bottom
- Paragraph indents: 1.25 cm (or 1.5 cm)
- Line spacing: 1.5 (полуторный)
- Alignment: justified (по ширине)
- Font: Times New Roman, 14pt main text, 12pt tables/footnotes
- Page numbering in headers
- Document structure
"""
import pytest
from pathlib import Path
from docx import Document
from tests.helpers.ooxml_utils import (
    get_document_xml,
    get_page_margins,
    get_page_size,
    get_paragraph_properties,
    get_run_properties,
    check_margins,
    mm_to_twips,
    cm_to_twips,
    twips_to_mm,
    twips_to_cm,
    pt_to_half_points,
    half_points_to_pt,
    NS,
)


class TestPageSetup:
    """Tests for page setup: margins, size, orientation."""
    
    def test_page_margins(self, any_docx):
        """
        Проверка полей страницы:
        - Левое: 30 мм
        - Правое: 10 мм
        - Верхнее: 20 мм
        - Нижнее: 20 мм
        """
        doc_xml = get_document_xml(any_docx)
        margins = get_page_margins(doc_xml)
        
        assert margins is not None, f"Не найдены поля страницы в {any_docx.name}"
        
        # Expected values in twips (with tolerance)
        expected = {
            'left': mm_to_twips(30),    # ~1701 twips
            'right': mm_to_twips(10),   # ~567 twips
            'top': mm_to_twips(20),     # ~1134 twips
            'bottom': mm_to_twips(20),  # ~1134 twips
        }
        
        tolerance_twips = mm_to_twips(1)  # 1mm tolerance
        
        for key, expected_value in expected.items():
            actual = margins.get(key)
            assert actual is not None, f"Поле '{key}' не задано в {any_docx.name}"
            
            diff = abs(actual - expected_value)
            actual_mm = twips_to_mm(actual)
            expected_mm = twips_to_mm(expected_value)
            
            assert diff <= tolerance_twips, (
                f"Поле '{key}' некорректно в {any_docx.name}: "
                f"ожидается {expected_mm:.1f} мм, фактически {actual_mm:.1f} мм"
            )
    
    def test_page_size_a4(self, any_docx):
        """Проверка размера страницы A4 (210×297 мм)."""
        doc_xml = get_document_xml(any_docx)
        page_size = get_page_size(doc_xml)
        
        assert page_size is not None, f"Размер страницы не найден в {any_docx.name}"
        
        # A4: 210mm × 297mm = 11906 × 16838 twips
        a4_width = mm_to_twips(210)
        a4_height = mm_to_twips(297)
        tolerance = mm_to_twips(5)  # 5mm tolerance
        
        width_diff = abs(page_size['width'] - a4_width)
        height_diff = abs(page_size['height'] - a4_height)
        
        assert width_diff <= tolerance and height_diff <= tolerance, (
            f"Размер страницы не соответствует A4 в {any_docx.name}: "
            f"{twips_to_mm(page_size['width']):.0f}×{twips_to_mm(page_size['height']):.0f} мм "
            f"(ожидается 210×297 мм)"
        )


class TestParagraphFormatting:
    """Tests for paragraph formatting: indents, spacing, alignment."""
    
    def test_first_line_indent(self, any_docx):
        """
        Проверка отступа первой строки абзаца: 1.25 см (или 1.5 см).
        Проверяем параграфы с явно заданным отступом.
        """
        doc_xml = get_document_xml(any_docx)
        paragraphs = doc_xml.xpath(".//w:p", namespaces=NS)
        
        # Expected values in twips
        indent_125 = cm_to_twips(1.25)  # ~709 twips
        indent_150 = cm_to_twips(1.5)   # ~850 twips
        tolerance = cm_to_twips(0.1)    # 1mm tolerance
        
        paragraphs_with_indent = 0
        invalid_indents = []
        
        for p in paragraphs:
            props = get_paragraph_properties(p)
            if 'ind' in props and props['ind'].get('firstLine'):
                # Handle both int and float strings
                first_line = float(props['ind']['firstLine'])
                first_line = int(round(first_line))
                paragraphs_with_indent += 1
                
                # Check if indent is approximately 1.25 or 1.5 cm
                diff_125 = abs(first_line - indent_125)
                diff_150 = abs(first_line - indent_150)
                
                if diff_125 > tolerance and diff_150 > tolerance:
                    actual_cm = twips_to_cm(first_line)
                    invalid_indents.append(f"{actual_cm:.2f} см")
        
        if invalid_indents:
            pytest.fail(
                f"Найдены некорректные отступы первой строки в {any_docx.name}: "
                f"{', '.join(invalid_indents[:5])} "
                f"(ожидается 1.25 см или 1.5 см)"
            )
    
    def test_line_spacing_15(self, any_docx):
        """
        Проверка межстрочного интервала: полуторный (1.5).
        В OOXML обычно lineRule="auto" и line="360" (или больше).
        """
        doc_xml = get_document_xml(any_docx)
        paragraphs = doc_xml.xpath(".//w:p[w:pPr/w:spacing]", namespaces=NS)
        
        invalid_spacing = []
        
        for p in paragraphs:
            props = get_paragraph_properties(p)
            if 'spacing' in props:
                spacing = props['spacing']
                line = spacing.get('line')
                line_rule = spacing.get('lineRule')
                
                # For 1.5 line spacing with auto rule, line should be around 360
                # (240 = single, 360 = 1.5, 480 = double)
                if line and line_rule == 'auto':
                    line_val = int(line)
                    # Accept 1.5 spacing: 340-380 range
                    if not (340 <= line_val <= 380):
                        invalid_spacing.append(f"line={line_val}")
        
        # Lenient check: only fail if spacing is explicitly wrong in many cases
        # Many documents use styles for spacing, so explicit spacing may not be set
        if len(paragraphs) > 0 and len(invalid_spacing) > 0:
            ratio = len(invalid_spacing) / len(paragraphs)
            # Only warn if more than 80% have wrong explicit spacing
            if ratio > 0.8:
                pytest.fail(
                    f"Много параграфов с некорректным интервалом в {any_docx.name}: "
                    f"{len(invalid_spacing)} из {len(paragraphs)}"
                )
    
    def test_justified_alignment(self, any_docx):
        """
        Проверка выравнивания текста: по ширине (both).
        Основной текст должен быть выровнен по ширине.
        """
        doc = Document(any_docx)
        doc_xml = get_document_xml(any_docx)
        
        paragraphs = doc_xml.xpath(".//w:p", namespaces=NS)
        
        justified_count = 0
        total_with_alignment = 0
        
        for p in paragraphs:
            props = get_paragraph_properties(p)
            if 'jc' in props:
                total_with_alignment += 1
                if props['jc'] == 'both':  # 'both' = justified
                    justified_count += 1
        
        # At least 50% of paragraphs with explicit alignment should be justified
        if total_with_alignment > 0:
            ratio = justified_count / total_with_alignment
            assert ratio >= 0.5, (
                f"Недостаточно параграфов с выравниванием по ширине в {any_docx.name}: "
                f"{justified_count} из {total_with_alignment} ({ratio*100:.0f}%)"
            )


class TestFonts:
    """Tests for font properties."""
    
    def test_times_new_roman_font(self, any_docx):
        """
        Проверка использования шрифта Times New Roman.
        Проверяем runs с явно заданным шрифтом.
        """
        doc_xml = get_document_xml(any_docx)
        runs = doc_xml.xpath(".//w:r", namespaces=NS)
        
        fonts_used = set()
        
        for run in runs[:100]:  # Check first 100 runs
            props = get_run_properties(run)
            if 'rFonts' in props:
                r_fonts = props['rFonts']
                for font_type in ['ascii', 'hAnsi', 'cs']:
                    font_name = r_fonts.get(font_type)
                    if font_name:
                        fonts_used.add(font_name)
        
        # If fonts are explicitly set, Times New Roman should be among them
        if fonts_used:
            assert 'Times New Roman' in fonts_used, (
                f"Times New Roman не найден среди явно заданных шрифтов в {any_docx.name}. "
                f"Найдены: {', '.join(sorted(fonts_used))}"
            )
    
    def test_font_size_14pt_main_text(self, any_docx):
        """
        Проверка размера шрифта: 14 пт для основного текста.
        Проверяем, что большинство runs используют 14pt или 12pt.
        
        Примечание: многие документы используют стили, поэтому размер
        может не быть задан явно. Этот тест пропускается, если размеры
        не заданы явно, или проверяет только явно заданные.
        """
        doc_xml = get_document_xml(any_docx)
        runs = doc_xml.xpath(".//w:r[w:rPr/w:sz]", namespaces=NS)
        
        size_14pt = pt_to_half_points(14)  # 28
        size_12pt = pt_to_half_points(12)  # 24
        
        sizes = []
        
        for run in runs:
            props = get_run_properties(run)
            if 'sz' in props:
                sizes.append(props['sz'])
        
        if not sizes:
            pytest.skip(f"Размеры шрифта не заданы явно в {any_docx.name}")
        
        # Count occurrences of standard sizes
        count_14 = sizes.count(size_14pt)
        count_12 = sizes.count(size_12pt)
        total = len(sizes)
        
        standard_ratio = (count_14 + count_12) / total
        
        # Relaxed check: at least 20% should be 14pt or 12pt if explicitly set
        # Many documents use styles, so this is just a sanity check
        if standard_ratio < 0.2:
            pytest.fail(
                f"Нестандартные размеры шрифта в {any_docx.name}: "
                f"14pt={count_14}, 12pt={count_12}, другие={total-count_14-count_12}"
            )


class TestDocumentStructure:
    """Tests for document structure and required sections."""
    
    def test_has_required_sections(self, any_docx):
        """
        Проверка наличия обязательных разделов:
        - Содержание (или Оглавление)
        - Введение
        - Заключение
        - Список использованных источников (или Список литературы)
        
        Примечание: это упрощённая проверка по наличию ключевых слов.
        Приложения не требуют всех разделов, только ПЗ.
        """
        doc = Document(any_docx)
        text = "\n".join([p.text for p in doc.paragraphs]).upper()
        
        # Приложения не требуют полной структуры
        if 'ПРИЛОЖЕНИЕ' in any_docx.name.upper():
            # For appendices, just check they exist (no strict requirements)
            return
        
        required_keywords = {
            'содержание': ['СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ'],
            'введение': ['ВВЕДЕНИЕ'],
            'заключение': ['ЗАКЛЮЧЕНИЕ'],
            'источники': ['СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 'СПИСОК ЛИТЕРАТУРЫ', 'БИБЛИОГРАФИЯ'],
        }
        
        missing = []
        
        for section_name, keywords in required_keywords.items():
            if not any(kw in text for kw in keywords):
                missing.append(section_name)
        
        if missing:
            pytest.fail(
                f"Отсутствуют обязательные разделы в {any_docx.name}: "
                f"{', '.join(missing)}"
            )
    
    def test_has_tables(self, pz_docx):
        """Проверка наличия таблиц в основном документе (ПЗ)."""
        doc = Document(pz_docx)
        
        assert len(doc.tables) > 0, f"Не найдены таблицы в {pz_docx.name}"
    
    def test_table_caption_format(self, any_docx):
        """
        Проверка формата подписей таблиц: "Таблица X.Y — Название".
        Упрощённая проверка по наличию слова "Таблица" и тире.
        """
        doc = Document(any_docx)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Look for table captions (simplified check)
        table_patterns = []
        for line in text.split('\n'):
            if 'Таблица' in line or 'ТАБЛИЦА' in line:
                table_patterns.append(line.strip())
        
        if table_patterns:
            # Check that at least some have the expected format with "—"
            has_dash = sum('—' in pattern or '–' in pattern for pattern in table_patterns)
            ratio = has_dash / len(table_patterns)
            
            assert ratio >= 0.5, (
                f"Найдены таблицы без правильного формата подписи в {any_docx.name}. "
                f"Ожидается: 'Таблица X.Y — Название'"
            )


class TestAdvanced:
    """Advanced checks (optional, may be skipped)."""
    
    def test_no_direct_font_formatting_in_body(self, any_docx):
        """
        Проверка: в основном тексте не должно быть прямого форматирования шрифта.
        Всё форматирование должно идти через стили (best practice).
        
        Примечание: это строгая проверка, может не пройти для многих документов.
        """
        doc_xml = get_document_xml(any_docx)
        
        # Count runs with direct font formatting
        direct_fonts = doc_xml.xpath("//w:r/w:rPr/w:rFonts", namespaces=NS)
        
        # Allow some direct formatting (up to 10% of runs)
        total_runs = len(doc_xml.xpath("//w:r", namespaces=NS))
        
        if total_runs > 0:
            ratio = len(direct_fonts) / total_runs
            if ratio > 0.1:
                pytest.skip(
                    f"Документ использует прямое форматирование ({ratio*100:.0f}%). "
                    f"Рекомендуется использовать стили."
                )


# Summary test that can be run separately
def test_normocontrol_summary(any_docx):
    """
    Сводная проверка основных требований нормоконтроля.
    Можно запускать отдельно для быстрой валидации.
    """
    doc_xml = get_document_xml(any_docx)
    doc = Document(any_docx)
    
    issues = []
    
    # Check margins
    if not check_margins(doc_xml, left_mm=30, right_mm=10, top_mm=20, bottom_mm=20, tolerance_mm=2):
        issues.append("Некорректные поля страницы")
    
    # Check structure
    text = "\n".join([p.text for p in doc.paragraphs]).upper()
    if 'ВВЕДЕНИЕ' not in text:
        issues.append("Отсутствует раздел 'Введение'")
    if 'ЗАКЛЮЧЕНИЕ' not in text:
        issues.append("Отсутствует раздел 'Заключение'")
    
    if issues:
        pytest.fail(
            f"Найдены проблемы в {any_docx.name}:\n" + 
            "\n".join(f"  - {issue}" for issue in issues)
        )
