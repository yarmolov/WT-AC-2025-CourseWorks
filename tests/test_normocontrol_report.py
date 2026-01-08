"""
Normocontrol tests with report generation (non-failing mode).

These tests collect all issues into a report instead of failing immediately.
Run with: pytest tests/test_normocontrol_report.py --report-format=markdown
"""
import pytest
from pathlib import Path
from docx import Document
from tests.helpers.ooxml_utils import (
    get_document_xml,
    get_page_margins,
    get_page_size,
    get_paragraph_properties,
    get_run_properties,
    mm_to_twips,
    cm_to_twips,
    twips_to_mm,
    twips_to_cm,
    pt_to_half_points,
    half_points_to_pt,
    find_paragraph_index,
    get_paragraph_text_preview,
    NS,
)


def test_all_documents_normocontrol(any_docx, normocontrol_report):
    """
    Comprehensive normocontrol check that collects all issues.
    This test never fails - it only collects issues into the report.
    """
    doc_name = any_docx.name
    normocontrol_report.add_document(doc_name)
    
    doc_xml = get_document_xml(any_docx)
    doc = Document(any_docx)
    
    # Check page margins
    _check_page_margins(any_docx, doc_xml, normocontrol_report)
    
    # Check page size
    _check_page_size(any_docx, doc_xml, normocontrol_report)
    
    # Check paragraph formatting
    _check_paragraph_indents(any_docx, doc_xml, normocontrol_report)
    _check_line_spacing(any_docx, doc_xml, normocontrol_report)
    _check_alignment(any_docx, doc_xml, normocontrol_report)
    
    # Check fonts
    _check_fonts(any_docx, doc_xml, normocontrol_report)
    _check_font_sizes(any_docx, doc_xml, normocontrol_report)
    
    # Check structure
    _check_document_structure(any_docx, doc, normocontrol_report)
    _check_table_captions(any_docx, doc, normocontrol_report)


def _check_page_margins(docx_path, doc_xml, report):
    """Check page margins."""
    doc_name = docx_path.name
    margins = get_page_margins(doc_xml)
    
    if not margins:
        report.add_issue(
            doc_name, "margins", "error",
            "Поля страницы не найдены",
            expected="Поля должны быть заданы",
            actual="Поля отсутствуют",
            location="Настройки страницы"
        )
        return
    
    expected_margins = {
        'left': (30, mm_to_twips(30)),
        'right': (10, mm_to_twips(10)),
        'top': (20, mm_to_twips(20)),
        'bottom': (20, mm_to_twips(20)),
    }
    
    tolerance_mm = 2
    tolerance_twips = mm_to_twips(tolerance_mm)
    
    for key, (expected_mm, expected_twips) in expected_margins.items():
        if key not in margins:
            report.add_issue(
                doc_name, "margins", "error",
                f"Поле '{key}' не задано",
                expected=f"{expected_mm} мм",
                actual="не задано",
                location="Разметка страницы → Поля"
            )
            continue
        
        actual_twips = margins[key]
        actual_mm = twips_to_mm(actual_twips)
        diff = abs(actual_twips - expected_twips)
        
        if diff > tolerance_twips:
            severity = "error" if diff > tolerance_twips * 2 else "warning"
            report.add_issue(
                doc_name, "margins", severity,
                f"Некорректное поле '{key}'",
                expected=f"{expected_mm} мм",
                actual=f"{actual_mm:.1f} мм",
                location="Разметка страницы → Поля → Настраиваемые поля"
            )


def _check_page_size(docx_path, doc_xml, report):
    """Check page size (A4)."""
    doc_name = docx_path.name
    page_size = get_page_size(doc_xml)
    
    if not page_size:
        report.add_issue(
            doc_name, "page_size", "warning",
            "Размер страницы не найден"
        )
        return
    
    a4_width = mm_to_twips(210)
    a4_height = mm_to_twips(297)
    tolerance = mm_to_twips(5)
    
    width_diff = abs(page_size['width'] - a4_width)
    height_diff = abs(page_size['height'] - a4_height)
    
    if width_diff > tolerance or height_diff > tolerance:
        actual_width_mm = twips_to_mm(page_size['width'])
        actual_height_mm = twips_to_mm(page_size['height'])
        report.add_issue(
            doc_name, "page_size", "warning",
            "Размер страницы не соответствует A4",
            expected="210×297 мм (A4)",
            actual=f"{actual_width_mm:.0f}×{actual_height_mm:.0f} мм"
        )


def _check_paragraph_indents(docx_path, doc_xml, report):
    """Check first-line indents."""
    doc_name = docx_path.name
    paragraphs = doc_xml.xpath(".//w:p", namespaces=NS)
    
    indent_125 = cm_to_twips(1.25)
    indent_150 = cm_to_twips(1.5)
    tolerance = cm_to_twips(0.1)
    
    invalid_indents = []
    problem_locations = []
    
    for idx, p in enumerate(paragraphs):
        props = get_paragraph_properties(p)
        if 'ind' in props and props['ind'].get('firstLine'):
            try:
                first_line = float(props['ind']['firstLine'])
                first_line = int(round(first_line))
                
                diff_125 = abs(first_line - indent_125)
                diff_150 = abs(first_line - indent_150)
                
                if diff_125 > tolerance and diff_150 > tolerance:
                    actual_cm = twips_to_cm(first_line)
                    if len(invalid_indents) < 10:  # Limit collected examples
                        invalid_indents.append(f"{actual_cm:.2f} см")
                        # Get paragraph preview for location
                        preview = get_paragraph_text_preview(p, 40)
                        problem_locations.append(f"Параграф {idx + 1}: '{preview}'")
            except (ValueError, TypeError):
                pass
    
    if invalid_indents:
        location = "; ".join(problem_locations[:3])  # Show first 3 locations
        if len(problem_locations) > 3:
            location += f" (и ещё {len(problem_locations) - 3})"
        
        report.add_issue(
            doc_name, "indents", "warning",
            f"Найдены некорректные отступы первой строки ({len(invalid_indents)} шт.)",
            expected="1.25 см или 1.5 см",
            actual=", ".join(invalid_indents[:5]),
            location=location
        )


def _check_line_spacing(docx_path, doc_xml, report):
    """Check line spacing."""
    doc_name = docx_path.name
    paragraphs = doc_xml.xpath(".//w:p[w:pPr/w:spacing]", namespaces=NS)
    
    invalid_count = 0
    first_problem_para = None
    
    for idx, p in enumerate(paragraphs):
        props = get_paragraph_properties(p)
        if 'spacing' in props:
            spacing = props['spacing']
            line = spacing.get('line')
            line_rule = spacing.get('lineRule')
            
            if line and line_rule == 'auto':
                try:
                    line_val = int(line)
                    if not (340 <= line_val <= 380):
                        invalid_count += 1
                        if first_problem_para is None:
                            first_problem_para = idx + 1
                except (ValueError, TypeError):
                    pass
    
    if len(paragraphs) > 0:
        ratio = invalid_count / len(paragraphs)
        if ratio > 0.8:
            location = f"Начиная с параграфа {first_problem_para}" if first_problem_para else "Весь документ"
            report.add_issue(
                doc_name, "spacing", "warning",
                f"Много параграфов с некорректным интервалом",
                expected="1.5 (полуторный)",
                actual=f"{invalid_count} из {len(paragraphs)} параграфов",
                location=location
            )


def _check_alignment(docx_path, doc_xml, report):
    """Check text alignment."""
    doc_name = docx_path.name
    paragraphs = doc_xml.xpath(".//w:p", namespaces=NS)
    
    justified_count = 0
    total_with_alignment = 0
    
    for p in paragraphs:
        props = get_paragraph_properties(p)
        if 'jc' in props:
            total_with_alignment += 1
            if props['jc'] == 'both':
                justified_count += 1
    
    if total_with_alignment > 0:
        ratio = justified_count / total_with_alignment
        if ratio < 0.5:
            report.add_issue(
                doc_name, "alignment", "warning",
                "Недостаточно параграфов с выравниванием по ширине",
                expected="Большинство параграфов по ширине",
                actual=f"{justified_count} из {total_with_alignment} ({ratio*100:.0f}%)"
            )


def _check_fonts(docx_path, doc_xml, report):
    """Check font usage."""
    doc_name = docx_path.name
    runs = doc_xml.xpath(".//w:r", namespaces=NS)
    
    fonts_used = set()
    
    for run in runs[:100]:
        props = get_run_properties(run)
        if 'rFonts' in props:
            r_fonts = props['rFonts']
            for font_type in ['ascii', 'hAnsi', 'cs']:
                font_name = r_fonts.get(font_type)
                if font_name:
                    fonts_used.add(font_name)
    
    if fonts_used and 'Times New Roman' not in fonts_used:
        report.add_issue(
            doc_name, "fonts", "info",
            "Times New Roman не найден среди явно заданных шрифтов",
            expected="Times New Roman",
            actual=", ".join(sorted(fonts_used)[:3])
        )


def _check_font_sizes(docx_path, doc_xml, report):
    """Check font sizes."""
    doc_name = docx_path.name
    runs = doc_xml.xpath(".//w:r[w:rPr/w:sz]", namespaces=NS)
    
    size_14pt = pt_to_half_points(14)
    size_12pt = pt_to_half_points(12)
    
    sizes = []
    
    for run in runs:
        props = get_run_properties(run)
        if 'sz' in props:
            sizes.append(props['sz'])
    
    if sizes:
        count_14 = sizes.count(size_14pt)
        count_12 = sizes.count(size_12pt)
        total = len(sizes)
        standard_ratio = (count_14 + count_12) / total
        
        if standard_ratio < 0.2:
            report.add_issue(
                doc_name, "fonts", "info",
                "Нестандартные размеры шрифта",
                expected="14pt (основной) или 12pt (таблицы)",
                actual=f"14pt={count_14}, 12pt={count_12}, другие={total-count_14-count_12}"
            )


def _check_document_structure(docx_path, doc, report):
    """Check document structure."""
    doc_name = docx_path.name
    
    # Приложения не требуют полной структуры
    if 'ПРИЛОЖЕНИЕ' in doc_name.upper():
        return
    
    text = "\n".join([p.text for p in doc.paragraphs]).upper()
    
    required_keywords = {
        'содержание': (['СОДЕРЖАНИЕ', 'ОГЛАВЛЕНИЕ'], 'Начало документа'),
        'введение': (['ВВЕДЕНИЕ'], 'После содержания'),
        'заключение': (['ЗАКЛЮЧЕНИЕ'], 'Конец основной части'),
        'источники': (['СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', 'СПИСОК ЛИТЕРАТУРЫ', 'БИБЛИОГРАФИЯ'], 'После заключения'),
    }
    
    for section_name, (keywords, location) in required_keywords.items():
        if not any(kw in text for kw in keywords):
            report.add_issue(
                doc_name, "structure", "error",
                f"Отсутствует раздел '{section_name}'",
                expected=f"Обязательный раздел: {keywords[0]}",
                actual="Раздел не найден",
                location=location
            )


def _check_table_captions(docx_path, doc, report):
    """Check table caption format."""
    doc_name = docx_path.name
    text = "\n".join([p.text for p in doc.paragraphs])
    
    table_patterns = []
    table_line_numbers = []
    
    for idx, line in enumerate(text.split('\n')):
        if 'Таблица' in line or 'ТАБЛИЦА' in line:
            table_patterns.append(line.strip())
            table_line_numbers.append(idx + 1)
    
    if table_patterns:
        has_dash = sum('—' in pattern or '–' in pattern for pattern in table_patterns)
        ratio = has_dash / len(table_patterns)
        
        if ratio < 0.5:
            # Find first problematic table
            first_bad = next((i for i, p in enumerate(table_patterns) if '—' not in p and '–' not in p), 0)
            location = f"Строка ~{table_line_numbers[first_bad]}" if first_bad < len(table_line_numbers) else "По всему документу"
            
            report.add_issue(
                doc_name, "tables", "warning",
                "Таблицы без правильного формата подписи",
                expected="Таблица X.Y — Название",
                actual=f"{len(table_patterns) - has_dash} из {len(table_patterns)} без тире",
                location=location
            )
